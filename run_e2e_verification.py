#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CMC End-to-End Full-Pipeline Verification
===========================================
Replicates the MATLAB pipeline in Python for two long texts:
  1. 繁体版出師表.docx  (Traditional Chinese)
  2. 背影.docx          (Simplified Chinese)

Outputs all 6 metrics:
  - Storage density (bits/nt)
  - GC content (%)
  - Maximum homopolymer length
  - Undesired motif content (%)
  - Hamming distance (minimum pairwise)
  - Minimum free energy (kcal/mol)
"""

import os, sys, csv, math, json, io
from pathlib import Path

# Fix Windows console encoding
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

# ─── Dependencies ───
import brotli
import docx
import opencc

PROJECT_DIR = Path(r"E:\研\李孜淇\Projects\CMC")
DATA_RAW    = PROJECT_DIR / "data_raw"
CHAR_CSV    = PROJECT_DIR / "character.csv"
OUTPUT_DIR  = PROJECT_DIR / "metrics"
OUTPUT_DIR.mkdir(exist_ok=True)

# ═══════════════════════════════════════════════════════════
# 1. Wubi Table Loading
# ═══════════════════════════════════════════════════════════

def load_wubi_table(csv_path):
    """Load character.csv → char_to_code, code_to_chars maps."""
    char_to_code = {}
    code_to_chars = {}   # code → ordered list of chars (preserving CSV order)

    with open(csv_path, 'r', encoding='utf-8') as f:
        reader = csv.reader(f)
        header = next(reader)
        for row in reader:
            if not row:
                continue
            code = row[0].strip()
            if not code:
                continue
            for col_idx in range(1, min(5, len(row))):
                ch = row[col_idx].strip()
                if ch and len(ch) == 1:
                    char_to_code[ch] = code
                    if code not in code_to_chars:
                        code_to_chars[code] = []
                    if ch not in code_to_chars[code]:
                        code_to_chars[code].append(ch)

    print(f"  Built lookup map with {len(char_to_code)} characters.")
    return char_to_code, code_to_chars


# ═══════════════════════════════════════════════════════════
# 2. Wubi Encoder
# ═══════════════════════════════════════════════════════════

def wubi_encode(text, char_to_code, code_to_chars):
    """Encode simplified Chinese text to wubi string + side-info."""
    wubi_parts = []
    punct_positions = []
    punct_chars = []
    collision_positions = []
    collision_indices = []
    char_idx = 0

    for ch in text:
        if ch in char_to_code:
            code = char_to_code[ch]
            wubi_parts.append(code)
            char_idx += 1

            # Collision detection
            if code in code_to_chars:
                candidates = code_to_chars[code]
                if len(candidates) > 1:
                    try:
                        idx = candidates.index(ch) + 1  # 1-based
                    except ValueError:
                        idx = 1
                    collision_positions.append(char_idx)
                    collision_indices.append(idx)
        else:
            punct_positions.append(char_idx)
            punct_chars.append(ch)

    wubi_str = '6'.join(wubi_parts)
    print(f"  Wubi length: {len(wubi_str)}, punctuation: {len(punct_positions)}, "
          f"collisions: {len(collision_positions)}")
    return wubi_str, punct_positions, punct_chars, collision_positions, collision_indices


# ═══════════════════════════════════════════════════════════
# 3. Wubi Decoder
# ═══════════════════════════════════════════════════════════

def wubi_decode(wubi_str, code_to_chars, punct_positions, punct_chars,
                collision_positions, collision_indices):
    """Decode wubi string back to Chinese text."""
    # Build collision lookup: char_idx → 1-based index
    collision_lookup = dict(zip(collision_positions, collision_indices))

    segments = wubi_str.split('6')
    chars = []
    for k, seg in enumerate(segments, start=1):
        if seg in code_to_chars:
            candidates = code_to_chars[seg]
            if len(candidates) > 1 and k in collision_lookup:
                idx = collision_lookup[k] - 1  # 0-based
                if 0 <= idx < len(candidates):
                    chars.append(candidates[idx])
                else:
                    chars.append(candidates[0])
            else:
                chars.append(candidates[0])

    # Re-insert punctuation
    result = []
    punct_idx = 0
    for i, ch in enumerate(chars, start=1):
        while punct_idx < len(punct_positions) and punct_positions[punct_idx] < i:
            result.append(punct_chars[punct_idx])
            punct_idx += 1
        result.append(ch)
    while punct_idx < len(punct_positions):
        result.append(punct_chars[punct_idx])
        punct_idx += 1

    return ''.join(result)


# ═══════════════════════════════════════════════════════════
# 4. Constrained DNA Encoder (get_allowed_bases + GC reorder)
# ═══════════════════════════════════════════════════════════

BAD_MOTIFS = ['TGC', 'CGC', 'GTC', 'GTG', 'GAC', 'CAC', 'GCG']

def build_allowed_lookup():
    """Pre-compute allowed bases for every 2-mer context."""
    bases = 'ACGT'
    lookup = {}
    for b1 in bases:
        for b2 in bases:
            ctx = b1 + b2
            forbidden = set()
            for mot in BAD_MOTIFS:
                if ctx[0] == mot[0] and ctx[1] == mot[1]:
                    forbidden.add(mot[2])
            # Homopolymer: if last 2 same, forbid that base
            if b1 == b2:
                forbidden.add(b1)
            lookup[ctx] = ''.join(b for b in bases if b not in forbidden)
    return lookup

ALLOWED_LOOKUP = build_allowed_lookup()

def reorder_bases_gc(allowed, gc_count, at_count):
    """Dynamically reorder allowed bases to balance GC ~50%."""
    total = gc_count + at_count
    if total == 0:
        return allowed
    gc_ratio = gc_count / total
    gc_bases = ''.join(b for b in allowed if b in 'GC')
    at_bases = ''.join(b for b in allowed if b in 'AT')
    if gc_ratio > 0.5:
        return at_bases + gc_bases
    else:
        return gc_bases + at_bases


def CMC_brotli_encode(wubi_str):
    """Encode wubi string → DNA sequence + meta, replicating MATLAB logic."""
    # Step 1: Brotli compression
    raw_bytes = wubi_str.encode('ascii')
    compressed = brotli.compress(raw_bytes, quality=11)
    comp_bytes = list(compressed)
    print(f"  Brotli(q11): {len(raw_bytes)} -> {len(comp_bytes)} bytes "
          f"({len(comp_bytes)/len(raw_bytes)*100:.1f}%)")

    # Step 2: Bytes → bitstream
    nbits = len(comp_bytes) * 8
    bitstream = []
    for byte_val in comp_bytes:
        for j in range(7, -1, -1):
            bitstream.append((byte_val >> j) & 1)
    assert len(bitstream) == nbits

    # Step 3: Constrained DNA encoding with limited ternary grouping
    max_ternary_groups = 25
    dna_buf = []
    bitpos = 0
    gc_count = 0
    at_count = 0

    ternary_buf = []
    ternary_pos = 0
    ternary_len = 0
    ternary_groups_used = 0

    while bitpos < nbits or ternary_pos < ternary_len:
        dna_pos = len(dna_buf)
        if dna_pos < 2:
            allowed = 'ACGT'
        else:
            ctx = dna_buf[-2] + dna_buf[-1]
            allowed = ALLOWED_LOOKUP[ctx]

        allowed = reorder_bases_gc(allowed, gc_count, at_count)
        n = len(allowed)

        if n >= 4:
            if bitpos + 1 < nbits:
                idx = bitstream[bitpos] * 2 + bitstream[bitpos + 1]
                bitpos += 2
            elif bitpos < nbits:
                idx = bitstream[bitpos]
                bitpos += 1
            else:
                break
        elif n == 3:
            if ternary_pos < ternary_len:
                idx = ternary_buf[ternary_pos]
                ternary_pos += 1
            elif (ternary_groups_used < max_ternary_groups and
                  (nbits - bitpos) >= 30):
                # Start new ternary group: consume 11 bits → 7 ternary digits
                val = 0
                for _ in range(11):
                    val = val * 2 + bitstream[bitpos]
                    bitpos += 1
                ternary_buf = []
                tmp = val
                for _ in range(7):
                    ternary_buf.insert(0, tmp % 3)
                    tmp //= 3
                ternary_len = 7
                ternary_pos = 0
                idx = ternary_buf[ternary_pos]
                ternary_pos += 1
                ternary_groups_used += 1
            elif bitpos < nbits:
                idx = bitstream[bitpos]
                bitpos += 1
            else:
                break
        else:  # n == 2
            if bitpos < nbits:
                idx = bitstream[bitpos]
                bitpos += 1
            else:
                break

        base = allowed[idx]
        dna_buf.append(base)
        if base in 'GC':
            gc_count += 1
        else:
            at_count += 1

    dna_sequence = ''.join(dna_buf)
    gc_ratio = gc_count / (gc_count + at_count) * 100 if (gc_count + at_count) > 0 else 0
    print(f"  Encoded: {len(dna_sequence)} bases, GC: {gc_ratio:.1f}%, "
          f"Ternary groups: {ternary_groups_used}/{max_ternary_groups}")

    meta = {
        'method': 'brotli_constrained_v4_limited_ternary',
        'originalLen': len(wubi_str),
        'compressedLen': len(comp_bytes),
        'totalBits': nbits,
        'gcRatio': gc_ratio,
        'maxTernaryGroups': max_ternary_groups,
        'ternaryGroupsUsed': ternary_groups_used,
    }
    return dna_sequence, meta


# ═══════════════════════════════════════════════════════════
# 5. Constrained DNA Decoder
# ═══════════════════════════════════════════════════════════

def CMC_brotli_decode(dna_sequence, meta):
    """Decode DNA → wubi string, replicating MATLAB logic."""
    dna_len = len(dna_sequence)
    nbits = meta['totalBits']
    max_ternary_groups = meta['maxTernaryGroups']

    bitstream = [0] * nbits
    gc_count = 0
    at_count = 0
    bits_consumed = 0
    ternary_groups_used = 0
    ternary_buf = []
    ternary_write_pos = 0

    for i in range(dna_len):
        if i < 2:
            allowed = 'ACGT'
        else:
            ctx = dna_sequence[i-2] + dna_sequence[i-1]
            allowed = ALLOWED_LOOKUP[ctx]

        allowed = reorder_bases_gc(allowed, gc_count, at_count)
        n = len(allowed)
        idx = allowed.index(dna_sequence[i])

        if n >= 4:
            bits_remaining = nbits - bits_consumed
            if bits_remaining >= 2:
                bitstream[bits_consumed]     = idx // 2
                bitstream[bits_consumed + 1] = idx % 2
                bits_consumed += 2
            elif bits_remaining == 1:
                bitstream[bits_consumed] = idx
                bits_consumed += 1
        elif n == 3:
            if ternary_buf:
                ternary_buf.append(idx)
                if len(ternary_buf) == 7:
                    val = 0
                    for t in ternary_buf:
                        val = val * 3 + t
                    for b in range(11):
                        bitstream[ternary_write_pos + b] = (val >> (10 - b)) & 1
                    ternary_buf = []
            else:
                bits_remaining = nbits - bits_consumed
                if (ternary_groups_used < max_ternary_groups and
                    bits_remaining >= 30):
                    ternary_write_pos = bits_consumed
                    bits_consumed += 11
                    ternary_buf = [idx]
                    ternary_groups_used += 1
                elif bits_remaining >= 1:
                    bitstream[bits_consumed] = idx
                    bits_consumed += 1
        else:  # n == 2
            bits_remaining = nbits - bits_consumed
            if bits_remaining >= 1:
                bitstream[bits_consumed] = idx
                bits_consumed += 1

        if dna_sequence[i] in 'GC':
            gc_count += 1
        else:
            at_count += 1

    # Handle remaining ternary digits
    if ternary_buf:
        nrem = len(ternary_buf)
        val = 0
        for t in ternary_buf:
            val = val * 3 + t
        val *= 3 ** (7 - nrem)
        for b in range(11):
            pos = ternary_write_pos + b
            if pos < nbits:
                bitstream[pos] = (val >> (10 - b)) & 1

    # Bitstream → bytes
    nbytes = meta['compressedLen']
    result_bytes = bytearray(nbytes)
    for i in range(nbytes):
        v = 0
        for j in range(8):
            bidx = i * 8 + j
            if bidx < len(bitstream):
                v = (v << 1) | bitstream[bidx]
            else:
                v <<= 1
        result_bytes[i] = v

    # Brotli decompress
    decompressed = brotli.decompress(bytes(result_bytes))
    decoded_str = decompressed.decode('ascii')
    print(f"  Decoded: {dna_len} bases -> {len(decoded_str)} symbols")
    return decoded_str


# ═══════════════════════════════════════════════════════════
# 6. Metric Computation
# ═══════════════════════════════════════════════════════════

def compute_storage_density(text, dna):
    """Erlich formula: CJK=16 bits, ASCII=8 bits."""
    info_bits = 0
    n_cjk = 0
    n_ascii = 0
    for ch in text:
        if ord(ch) > 127:
            info_bits += 16
            n_cjk += 1
        else:
            info_bits += 8
            n_ascii += 1
    density = info_bits / len(dna) if len(dna) > 0 else 0
    return density, info_bits, n_cjk, n_ascii


def compute_gc_content(dna):
    gc = sum(1 for b in dna if b in 'GC')
    return gc / len(dna) * 100 if dna else 0


def compute_max_homopolymer(dna):
    if not dna:
        return 0
    max_run = 1
    current_run = 1
    for i in range(1, len(dna)):
        if dna[i] == dna[i-1]:
            current_run += 1
        else:
            max_run = max(max_run, current_run)
            current_run = 1
    return max(max_run, current_run)


def compute_undesired_motifs(dna):
    total_triplets = len(dna) - 2
    motif_counts = {}
    total_bad = 0
    for motif in BAD_MOTIFS:
        count = 0
        for i in range(len(dna) - 2):
            if dna[i:i+3] == motif:
                count += 1
        motif_counts[motif] = count
        total_bad += count
    bad_ratio = total_bad / total_triplets * 100 if total_triplets > 0 else 0
    return total_bad, total_triplets, bad_ratio, motif_counts


def compute_hamming_distances(dna, frag_len=120):
    """Compute pairwise Hamming distances between fragments."""
    num_frags = min(6, len(dna) // frag_len)
    fragments = []
    for i in range(num_frags):
        start = i * frag_len
        end = start + frag_len
        fragments.append(dna[start:end])

    min_hamming = float('inf')
    pairs = []
    for i in range(num_frags):
        for j in range(i+1, num_frags):
            hd = sum(1 for a, b in zip(fragments[i], fragments[j]) if a != b)
            pairs.append((i+1, j+1, hd))
            if hd < min_hamming:
                min_hamming = hd

    return min_hamming, pairs, fragments, num_frags


def compute_free_energy_nn(dna, frag_len=120):
    """Nearest-neighbor model for MFE estimation (same as MATLAB fallback)."""
    nn_params = {
        'AA': -1.0, 'TT': -1.0,
        'AT': -0.88, 'TA': -0.58,
        'CA': -1.45, 'TG': -1.45,
        'GT': -1.44, 'AC': -1.44,
        'CT': -1.28, 'AG': -1.28,
        'GA': -1.30, 'TC': -1.30,
        'CG': -2.17, 'GC': -2.24,
        'GG': -1.84, 'CC': -1.84,
    }

    num_frags = min(6, len(dna) // frag_len)
    energies = []
    for i in range(num_frags):
        start = i * frag_len
        seq = dna[start:start+frag_len]
        dg = 0.0
        for j in range(len(seq) - 1):
            dinuc = seq[j:j+2]
            dg += nn_params.get(dinuc, -1.0)
        energies.append(dg)

    # Try ViennaRNA
    try:
        import RNA
        print("  Using ViennaRNA for MFE calculation...")
        energies_vienna = []
        for i in range(num_frags):
            start = i * frag_len
            seq = dna[start:start+frag_len]
            # DNA→RNA for ViennaRNA: T→U
            rna_seq = seq.replace('T', 'U')
            structure, mfe = RNA.fold(rna_seq)
            energies_vienna.append(mfe)
        return energies_vienna, num_frags, True
    except ImportError:
        print("  ViennaRNA not available, using nearest-neighbor approximation...")
        return energies, num_frags, False


# ═══════════════════════════════════════════════════════════
# 7. Main Pipeline
# ═══════════════════════════════════════════════════════════

def run_pipeline(docx_path, label, is_traditional=False):
    """Run the complete CMC pipeline for one document."""
    print(f"\n{'='*70}")
    print(f"  CMC End-to-End Verification: {label}")
    print(f"  File: {docx_path}")
    print(f"{'='*70}")

    # --- Step 1: Read docx ---
    print("\n[1] Reading .docx...")
    doc = docx.Document(str(docx_path))
    if is_traditional:
        # Traditional: join all paragraphs (matching MATLAB run_docx_test.m)
        text = ''.join(p.text for p in doc.paragraphs)
    else:
        # Simplified: join non-empty paragraphs (matching MATLAB run_docx_test_simp.m)
        text = ''.join(p.text for p in doc.paragraphs if p.text.strip())
    print(f"  Total characters: {len(text)}")
    print(f"  Preview: {text[:50]}...")

    # --- Step 2: Traditional → Simplified (if needed) ---
    if is_traditional:
        print("\n[2] Converting Traditional → Simplified...")
        converter_t2s = opencc.OpenCC('t2s')
        converter_s2t = opencc.OpenCC('s2t')
        text_simp = converter_t2s.convert(text)
        print(f"  Preview: {text_simp[:50]}...")
    else:
        text_simp = text
        converter_s2t = None
        print("\n[2] Already Simplified Chinese, no conversion needed.")

    # --- Step 3: Wubi encoding ---
    print("\n[3] Wubi encoding...")
    char_to_code, code_to_chars = load_wubi_table(CHAR_CSV)
    wubi_str, punct_pos, punct_ch, coll_pos, coll_idx = wubi_encode(
        text_simp, char_to_code, code_to_chars)
    print(f"  Preview: {wubi_str[:80]}...")

    # --- Step 4: CMC Brotli DNA encoding ---
    print("\n[4] CMC Brotli + constrained DNA encoding...")
    dna, meta = CMC_brotli_encode(wubi_str)
    print(f"  DNA length: {len(dna)} nt")
    print(f"  Preview: {dna[:80]}...")

    # --- Step 5: Decode and verify ---
    print("\n[5] Decoding and verification...")
    decoded_wubi = CMC_brotli_decode(dna, meta)
    decoded_simp = wubi_decode(decoded_wubi, code_to_chars,
                                punct_pos, punct_ch, coll_pos, coll_idx)

    wubi_match = (wubi_str == decoded_wubi)
    print(f"  Wubi match: {'OK ✓' if wubi_match else 'FAIL ✗'}")

    if is_traditional and converter_s2t:
        decoded_trad = converter_s2t.convert(decoded_simp)
        # OpenCC round-trip correction
        converter_t2s_fix = opencc.OpenCC('t2s')
        decoded_list = list(decoded_trad)
        for i in range(min(len(decoded_list), len(text))):
            if decoded_list[i] != text[i]:
                orig_s = converter_t2s_fix.convert(text[i])
                dec_s  = converter_t2s_fix.convert(decoded_list[i])
                if orig_s == dec_s:
                    decoded_list[i] = text[i]
        decoded_final = ''.join(decoded_list)
        text_match = (text == decoded_final)
    else:
        decoded_final = decoded_simp
        text_match = (text_simp == decoded_final)

    print(f"  Full-text match: {'OK ✓' if text_match else 'Partial differences'}")

    if not text_match:
        min_len = min(len(text if is_traditional else text_simp), len(decoded_final))
        ref = text if is_traditional else text_simp
        diffs = [i for i in range(min_len) if ref[i] != decoded_final[i]]
        if diffs:
            print(f"  Differences: {len(diffs)}")
            for d in diffs[:10]:
                print(f"    Position {d}: original[{ref[d]}] vs decoded[{decoded_final[d]}]")

    # --- Step 6: Compute all metrics ---
    print("\n[6] Computing metrics...")

    # 6a. Storage density
    density, info_bits, n_cjk, n_ascii = compute_storage_density(text, dna)
    print(f"\n  --- Storage Density ---")
    print(f"  Info bits: {info_bits} (CJK: {n_cjk}×16, ASCII: {n_ascii}×8)")
    print(f"  DNA length: {len(dna)} nt")
    print(f"  Storage density (Erlich): {density:.4f} bits/nt")

    # 6b. GC content
    gc_pct = compute_gc_content(dna)
    print(f"\n  --- GC Content ---")
    print(f"  GC content: {gc_pct:.2f}%")

    # 6c. Max homopolymer
    max_homo = compute_max_homopolymer(dna)
    print(f"\n  --- Maximum Homopolymer Length ---")
    print(f"  Max homopolymer: {max_homo} (target ≤ 2)")
    print(f"  {'✓ Constraint satisfied' if max_homo <= 2 else '✗ Constraint violated'}")

    # 6d. Undesired motifs
    total_bad, total_tri, bad_ratio, motif_counts = compute_undesired_motifs(dna)
    print(f"\n  --- Undesired Motif Content ---")
    for mot, cnt in motif_counts.items():
        print(f"    {mot}: {cnt}")
    print(f"  Total bad motifs: {total_bad} / {total_tri} triplets")
    print(f"  Bad motif ratio: {bad_ratio:.4f}% (target < 0.6%)")
    print(f"  {'✓ Constraint satisfied' if bad_ratio < 0.6 else '✗ Constraint violated'}")

    # 6e. Hamming distance
    min_hd, hd_pairs, fragments, num_frags = compute_hamming_distances(dna)
    print(f"\n  --- Hamming Distance ---")
    print(f"  Extracted {num_frags} fragments of 120 nt")
    for i, j, hd in hd_pairs:
        print(f"    Fragment {i} vs {j}: Hamming distance = {hd}")
    print(f"  Minimum Hamming distance: {min_hd} (target > 60)")
    print(f"  {'✓ Constraint satisfied' if min_hd > 60 else '✗ Constraint not met'}")

    # 6f. Minimum free energy
    energies, nf, used_vienna = compute_free_energy_nn(dna)
    method_label = "ViennaRNA" if used_vienna else "NN approximation"
    print(f"\n  --- Minimum Free Energy ({method_label}) ---")
    for i, e in enumerate(energies):
        print(f"    Fragment {i+1}: ΔG = {e:.1f} kcal/mol")
    avg_fe = sum(energies) / len(energies) if energies else 0
    min_fe = min(energies) if energies else 0
    print(f"  Average ΔG: {avg_fe:.1f} kcal/mol")
    print(f"  Minimum ΔG: {min_fe:.1f} kcal/mol (target > -30)")
    print(f"  {'✓ Constraint satisfied' if min_fe > -30 else '✗ Secondary structure risk'}")

    # --- Summary Report ---
    print(f"\n{'='*70}")
    print(f"  ╔══════════════════════════════════════════════╗")
    print(f"  ║   CMC Quality Report: {label:<22s} ║")
    print(f"  ╠══════════════════════════════════════════════╣")
    print(f"  ║ Original chars:       {len(text):>6d}              ║")
    print(f"  ║ DNA length:           {len(dna):>6d} nt           ║")
    print(f"  ║ Storage density:      {density:>6.4f} bits/nt     ║")
    print(f"  ║ GC content:           {gc_pct:>6.2f}%             ║")
    print(f"  ║ Max homopolymer:      {max_homo:>6d}              ║")
    print(f"  ║ Bad motif ratio:      {bad_ratio:>6.4f}%           ║")
    print(f"  ║ Avg free energy:      {avg_fe:>6.1f} kcal/mol    ║")
    print(f"  ║ Min free energy:      {min_fe:>6.1f} kcal/mol    ║")
    print(f"  ║ Min Hamming dist:     {min_hd:>6}              ║")
    print(f"  ║ Wubi round-trip:      {'PASS' if wubi_match else 'FAIL':>6s}              ║")
    print(f"  ║ Full-text round-trip: {'PASS' if text_match else 'FAIL':>6s}              ║")
    print(f"  ╚══════════════════════════════════════════════╝")

    return {
        'label': label,
        'chars': len(text),
        'dna_length': len(dna),
        'storage_density': density,
        'gc_content': gc_pct,
        'max_homopolymer': max_homo,
        'bad_motif_ratio': bad_ratio,
        'bad_motif_count': total_bad,
        'avg_free_energy': avg_fe,
        'min_free_energy': min_fe,
        'min_hamming': min_hd,
        'wubi_match': wubi_match,
        'text_match': text_match,
        'compressed_bytes': meta['compressedLen'],
        'wubi_length': meta['originalLen'],
        'info_bits': info_bits,
        'n_cjk': n_cjk,
        'n_ascii': n_ascii,
        'energies': energies,
        'hamming_pairs': hd_pairs,
        'motif_detail': motif_counts,
    }


# ═══════════════════════════════════════════════════════════
# 8. Entry Point
# ═══════════════════════════════════════════════════════════

if __name__ == '__main__':
    print("=" * 70)
    print("  CMC End-to-End Full-Pipeline Verification")
    print("  Date: 2026-03-26")
    print("=" * 70)

    results = []

    # Text 1: Traditional Chinese 出師表
    r1 = run_pipeline(
        DATA_RAW / "繁体版出師表.docx",
        "出師表(繁体)",
        is_traditional=True
    )
    results.append(r1)

    # Text 2: Simplified Chinese 背影
    r2 = run_pipeline(
        DATA_RAW / "背影.docx",
        "背影(简体)",
        is_traditional=False
    )
    results.append(r2)

    # ─── Combined Summary Table ───
    print(f"\n\n{'='*70}")
    print("  COMBINED COMPARISON TABLE")
    print(f"{'='*70}")
    header = f"{'Metric':<30s} | {'出師表(繁体)':>15s} | {'背影(简体)':>15s}"
    print(header)
    print("-" * len(header))
    rows = [
        ("Original chars",        f"{r1['chars']}",             f"{r2['chars']}"),
        ("DNA length (nt)",       f"{r1['dna_length']}",        f"{r2['dna_length']}"),
        ("Storage density (b/nt)",f"{r1['storage_density']:.4f}",f"{r2['storage_density']:.4f}"),
        ("GC content (%)",        f"{r1['gc_content']:.2f}",    f"{r2['gc_content']:.2f}"),
        ("Max homopolymer",       f"{r1['max_homopolymer']}",   f"{r2['max_homopolymer']}"),
        ("Bad motif ratio (%)",   f"{r1['bad_motif_ratio']:.4f}",f"{r2['bad_motif_ratio']:.4f}"),
        ("Avg ΔG (kcal/mol)",     f"{r1['avg_free_energy']:.1f}",f"{r2['avg_free_energy']:.1f}"),
        ("Min ΔG (kcal/mol)",     f"{r1['min_free_energy']:.1f}",f"{r2['min_free_energy']:.1f}"),
        ("Min Hamming distance",  f"{r1['min_hamming']}",       f"{r2['min_hamming']}"),
        ("Wubi round-trip",       "PASS" if r1['wubi_match'] else "FAIL",
                                  "PASS" if r2['wubi_match'] else "FAIL"),
        ("Text round-trip",       "PASS" if r1['text_match'] else "FAIL",
                                  "PASS" if r2['text_match'] else "FAIL"),
    ]
    for name, v1, v2 in rows:
        print(f"  {name:<28s} | {v1:>15s} | {v2:>15s}")

    # ─── Save to CSV ───
    csv_path = OUTPUT_DIR / "CMC_e2e_verification.csv"
    with open(csv_path, 'w', newline='', encoding='utf-8') as f:
        w = csv.writer(f)
        w.writerow(["Metric", "出師表(繁体)", "背影(简体)"])
        for name, v1, v2 in rows:
            w.writerow([name.strip(), v1.strip(), v2.strip()])
    print(f"\n  Results saved to: {csv_path}")

    print(f"\n{'='*70}")
    print("  Verification complete.")
    print(f"{'='*70}")

